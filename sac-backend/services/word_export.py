from io import BytesIO
from copy import deepcopy
from pathlib import Path
import unicodedata

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from models import SAC
from schemas import fmt_date

TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "templates" / "F01-PD-PR-02_05_SOLICITUD_Y_SEGUIMIENTO_DE_ACCION_CORRECTIVA_V05.docx"
FONT_NAME = "Arial"
FONT_SIZE = Pt(9)


def area_abbr(area: str) -> str:
    normalized = unicodedata.normalize("NFKD", area or "")
    text = "".join(ch for ch in normalized if not unicodedata.combining(ch)).upper()
    if "INVESTIG" in text:
        return "INV"
    known = {"CAP", "CID", "CIS", "SSOMA"}
    if text in known:
        return text
    words = [w for w in text.replace("&", " ").replace("-", " ").split() if w]
    skip = {"DE", "DEL", "LA", "LAS", "LOS", "Y", "EN", "AREA", "UNIDAD"}
    keywords = [w for w in words if w not in skip and w not in {"DIRECCION", "OFICINA", "PROGRAMA"}]
    if len(keywords) == 1:
        return keywords[0][:3]
    letters = [w[0] for w in keywords]
    return "".join(letters[:5]) or "AREA"


def replace_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    for key, value in replacements.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", str(value or ""))
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)


def sac_replacements(sac: SAC) -> dict:
    return {
        "codigo": sac.codigo,
        "code": sac.code,
        "campus": sac.campus,
        "proceso": sac.proceso,
        "proceso_sgc": sac.proceso_sgc,
        "norma": sac.norma,
        "clausula": sac.clausula,
        "originador": sac.originador,
        "fecha_registro": fmt_date(sac.fecha_registro),
        "fecha_compromiso": fmt_date(sac.fecha_compromiso),
        "responsable": sac.responsable,
        "prioridad": sac.prioridad,
        "fuente": sac.fuente,
        "nc": sac.nc,
        "accion_inmediata": sac.accion_inmediata,
        "accion_inm_responsable": sac.accion_inm_responsable,
        "accion_inm_fecha": fmt_date(sac.accion_inm_fecha),
        "analisis_causa": sac.analisis_causa,
        "estado": sac.estado,
        "implementacion": f"{sac.implementacion}%",
        "eficacia": sac.eficacia,
    }


def cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def format_run(run, bold=None) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def format_paragraph(paragraph, align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def format_document(document: Document) -> None:
    for paragraph in document.paragraphs:
        format_paragraph(paragraph)
        for run in paragraph.runs:
            format_run(run)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    format_paragraph(paragraph)
                    for run in paragraph.runs:
                        format_run(run)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text or "")
    for paragraph in cell.paragraphs:
        format_paragraph(paragraph)
        for run in paragraph.runs:
            format_run(run, bold=False)


def append_to_cell(cell, text: str, sep: str = " ") -> None:
    current = cell_text(cell)
    set_cell_text(cell, f"{current}{sep}{text}".strip() if current else text)


def set_label_value(cell, label: str, value: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(label)
    format_run(label_run, bold=True)
    if value:
        value_run = paragraph.add_run(f" {value}")
        format_run(value_run, bold=False)


def set_multiline_cell(cell, lines) -> None:
    cell.text = ""
    first = True
    for text, bold in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = paragraph.add_run(str(text or ""))
        format_run(run, bold=bold)


def add_plan_rows(table, count: int) -> None:
    if count <= 1:
        return
    previous = table.rows[17]._tr
    for _ in range(count - 1):
        new_row = deepcopy(table.rows[17]._tr)
        previous.addnext(new_row)
        previous = new_row


def remove_row_fixed_height(row) -> None:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return
    for node in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(node)


def set_action_cell(cell, text: str) -> None:
    set_cell_text(cell, text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=70, start=95, bottom=70, end=95)
    for paragraph in cell.paragraphs:
        format_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)


def source_options(fuente: str) -> str:
    selected = (fuente or "").strip().lower()
    options = [
        ("Auditorías internas o externas", {"auditoria interna", "auditoria externa"}),
        ("Evaluación del plan estratégico o plan operativo", {"evaluacion plan estrategico", "plan estrategico"}),
        ("Salidas no conformes", {"salidas no conformes"}),
        ("Resultados de las Revisiones por la Dirección", {"revision por direccion"}),
        ("Quejas y reclamos", {"quejas repetitivas"}),
        ("Satisfacción del cliente", {"satisfaccion del cliente"}),
        ("Medición y control", {"medicion y control"}),
        ("Otras fuentes", {"otras fuentes"}),
    ]
    lines = []
    for label, values in options:
        mark = "\u2612" if selected in values else "\u2610"
        lines.append(f"{mark} {label}")
    return "\n".join(lines)


def yes_no(value: str) -> str:
    value = (value or "").strip().lower()
    yes = "\u2612" if value == "si" else "\u2610"
    no = "\u2612" if value == "no" else "\u2610"
    return f"{yes} SI    {no} NO"


def fill_official_template(document: Document, sac: SAC) -> None:
    table = document.tables[0]

    display_code = sac.codigo or f"SAC-{sac.code}"
    set_label_value(table.cell(0, 0), "Solicitud N°", display_code)
    set_label_value(table.cell(1, 0), "Proceso:", sac.proceso_sgc or sac.proceso)
    set_cell_text(table.cell(3, 0), source_options(sac.fuente))

    set_multiline_cell(
        table.cell(5, 0),
        [
            (sac.procedimiento, True),
            (sac.nc or sac.descripcion, False),
        ],
    )
    set_cell_text(table.cell(6, 2), sac.clausula)
    set_cell_text(table.cell(7, 2), sac.norma)
    set_label_value(table.cell(8, 0), "Nombre del originador:", sac.originador)
    set_label_value(table.cell(8, 5), "Fecha:", fmt_date(sac.fecha_registro))

    set_cell_text(table.cell(10, 0), sac.accion_inmediata or "No aplica")
    set_label_value(table.cell(11, 0), "Responsable:", sac.accion_inm_responsable or "No aplica")
    set_label_value(table.cell(11, 4), "Fecha:", fmt_date(sac.accion_inm_fecha) or "No aplica")

    set_cell_text(table.cell(13, 0), sac.analisis_causa)

    actions = sorted(sac.plan_acciones, key=lambda x: x.orden)
    add_plan_rows(table, len(actions))
    for idx, action in enumerate(actions):
        row_idx = 17 + idx
        remove_row_fixed_height(table.rows[row_idx])
        set_action_cell(table.cell(row_idx, 0), str(action.orden))
        set_action_cell(table.cell(row_idx, 1), action.desc)
        set_action_cell(table.cell(row_idx, 3), action.responsable)
        set_action_cell(table.cell(row_idx, 6), fmt_date(action.fecha))

    after_actions = max(len(actions), 1) - 1
    set_label_value(table.cell(18 + after_actions, 0), "Acciones planteadas por:", sac.responsable or sac.originador)
    set_label_value(table.cell(18 + after_actions, 5), "Fecha:", fmt_date(sac.fecha_registro))

    set_cell_text(table.cell(20 + after_actions, 0), sac.verif_impl_desc)
    set_label_value(table.cell(21 + after_actions, 0), "Verificado por:", sac.verif_impl_por)
    set_label_value(table.cell(21 + after_actions, 5), "Fecha:", fmt_date(sac.verif_impl_fecha))
    set_label_value(table.cell(22 + after_actions, 0), "Verificar la eficacia a partir de:", sac.verif_impl_eficacia_desde)

    set_cell_text(
        table.cell(24 + after_actions, 0),
        "\n".join([
            "(Detallar los registros o documentación revisada)",
            sac.verif_efic_docs or "",
            "",
            f"¿La acción tomada fue eficaz? {yes_no(sac.verif_efic_eficaz)}",
            f"¿Se cierra la no conformidad? {yes_no(sac.verif_efic_cierra)}",
        ]).strip(),
    )
    set_label_value(table.cell(25 + after_actions, 0), "Verificado por:", sac.verif_efic_por)
    set_label_value(table.cell(25 + after_actions, 5), "Fecha:", fmt_date(sac.verif_efic_fecha))
    set_label_value(table.cell(26 + after_actions, 0), "OBSERVACIONES:", sac.verif_efic_obs)


def append_plan_table(document: Document, sac: SAC) -> None:
    document.add_heading("Plan de accion", level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["N", "Descripcion", "Responsable", "Fecha", "Estado"]
    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title
    for action in sorted(sac.plan_acciones, key=lambda x: x.orden):
        row = table.add_row().cells
        row[0].text = str(action.orden)
        row[1].text = action.desc
        row[2].text = action.responsable
        row[3].text = fmt_date(action.fecha)
        row[4].text = action.estado


def build_fallback_document(sac: SAC) -> Document:
    document = Document()
    document.add_heading("F01-PD-PR-02.05 - Solicitud de Accion Correctiva", level=1)
    document.add_paragraph(f"Codigo: {sac.codigo}")
    document.add_paragraph(f"Campus: {sac.campus}")
    document.add_paragraph(f"Area / proceso: {sac.proceso}")
    document.add_paragraph(f"Proceso SGC: {sac.proceso_sgc}")
    document.add_paragraph(f"Norma / clausula: {sac.norma} - {sac.clausula}")
    document.add_paragraph(f"Fuente: {sac.fuente}")
    document.add_paragraph(f"Prioridad: {sac.prioridad}")
    document.add_paragraph(f"Responsable: {sac.responsable}")
    document.add_paragraph(f"Fecha de registro: {fmt_date(sac.fecha_registro)}")
    document.add_heading("Descripcion de la no conformidad", level=2)
    document.add_paragraph(sac.nc or sac.descripcion)
    document.add_heading("Accion inmediata", level=2)
    document.add_paragraph(sac.accion_inmediata)
    document.add_heading("Analisis de causa raiz", level=2)
    document.add_paragraph(sac.analisis_causa)
    append_plan_table(document, sac)
    return document


def build_docx(sac: SAC) -> BytesIO:
    if TEMPLATE_PATH.exists():
        document = Document(TEMPLATE_PATH)
        replacements = sac_replacements(sac)
        for paragraph in document.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for table in document.tables:
            replace_in_table(table, replacements)
        if document.tables:
            fill_official_template(document, sac)
        format_document(document)
    else:
        document = build_fallback_document(sac)
        format_document(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
